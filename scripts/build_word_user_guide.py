from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Excel報名資料自動填寫工具_使用說明.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WARNING_FILL = "FFF4CE"
WARNING_BORDER = "D6A400"
BORDER = "B7C3D0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)

    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for side, value in values.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            if index >= len(row.cells):
                continue
            row.cells[index].width = width


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold=False, color=None, size=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_note_box(doc, title, lines, fill=WARNING_FILL, border=WARNING_BORDER):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [Inches(6.5)])
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, border, "8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    set_run_font(run, bold=True, color=DARK_BLUE)
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        set_run_font(r)
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_cell_margins(table)

    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True, color=DARK_BLUE)

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run)

    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Excel 報名資料自動填寫工具 使用說明")
    set_run_font(run, size=22, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("給一般使用者的快速操作指南")
    set_run_font(run, size=12, color="555555")

    add_note_box(
        doc,
        "文件重點",
        [
            "這份文件說明如何選擇 Excel、貼上報名/付款文字、確認預覽並寫入資料。",
            "正式使用前請先備份 Excel；如果 Windows 顯示不明來源提醒，請先確認程式來源可信再允許執行。",
        ],
        fill=LIGHT_GRAY,
        border=BORDER,
    )

    add_heading(doc, "1. 這個工具可以做什麼")
    add_paragraph(
        doc,
        "這個桌面工具會讀取你貼上的報名或付款文字，整理成可檢查的預覽資料，確認後自動新增到指定的 Excel .xlsx 檔案。使用者不需要手動一列一列輸入資料。"
    )
    add_table(
        doc,
        ["步驟", "動作", "結果"],
        [
            ["1", "準備 Excel 檔", "確認檔案格式為 .xlsx，並先備份正式資料。"],
            ["2", "貼上報名/付款文字", "工具自動解析姓名、課程、金額、付款日期與付款代碼。"],
            ["3", "檢查右側預覽", "確認筆數、金額、付款方式與日期是否正確。"],
            ["4", "按下寫入按鈕", "工具從第一個工作表第 3 列開始插入資料。"],
            ["5", "開啟 Excel 檢查", "確認新增資料列是否符合預期。"],
        ],
        [Inches(0.65), Inches(2.0), Inches(3.85)],
    )

    add_heading(doc, "2. 開始前檢查")
    add_table(
        doc,
        ["檢查項目", "請確認"],
        [
            ["Excel 檔案", "必須是 .xlsx 格式。"],
            ["工作表", "工具會寫入第一個工作表。"],
            ["寫入位置", "新資料會從第 3 列開始插入，原本第 3 列以下資料會往下移。"],
            ["資料備份", "正式資料第一次使用前，請先複製一份 Excel 備份。"],
            ["檔案狀態", "寫入前請關閉正在開啟的同一份 Excel 檔案。"],
        ],
        [Inches(1.7), Inches(4.8)],
    )

    add_heading(doc, "3. 第一次開啟時的不明來源提醒")
    add_note_box(
        doc,
        "重要提醒：Windows 可能要求信任或同意執行",
        [
            "如果這個工具是可攜版或內部自製工具，Windows 可能顯示「Windows 已保護您的電腦」、「不明的發行者」、「不明來源」或類似安全提醒。",
            "請先確認檔案來源是公司內部、專案負責人或可信任的分享位置，再選擇「其他資訊」、「仍要執行」、「允許」或「信任」。",
            "如果你不確定檔案來源，請不要直接同意，先詢問工具提供者或資訊人員。",
        ],
    )
    add_paragraph(
        doc,
        "這個提醒通常是因為程式尚未做正式程式碼簽章，並不一定代表程式有問題；但仍應只執行來源可信的版本。"
    )

    add_heading(doc, "4. 主畫面區塊")
    add_table(
        doc,
        ["畫面區塊", "用途", "使用者要做什麼"],
        [
            ["上方操作列", "顯示工具名稱與寫入按鈕。", "完成選檔與預覽確認後，按下寫入。"],
            ["左側資料輸入", "選擇 Excel 檔並貼上報名/付款文字。", "按「選擇」挑 .xlsx，將文字貼到大文字框。"],
            ["右側資料預覽", "顯示解析結果、每筆資料與處理訊息。", "確認姓名、金額、付款日期、付款代碼與筆數。"],
        ],
        [Inches(1.45), Inches(2.25), Inches(2.8)],
    )

    add_heading(doc, "5. 操作步驟")
    for text in [
        "按「選擇」挑選要寫入的 Excel .xlsx 檔案。",
        "把報名或付款文字貼到左側大文字框。",
        "檢查右側預覽，確認資料筆數、姓名、付款方式、付款日期、付款代碼與金額。",
        "如果畫面出現紅色錯誤提示，先回到左側補齊缺少的資料。",
        "確認無誤後按下上方寫入按鈕。",
        "寫入完成後，開啟 Excel 檢查第 3 列起新增的資料。",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "6. 建議貼上文字格式")
    add_paragraph(doc, "貼上的文字建議包含以下資訊。文字不一定要完全相同，但資料越清楚，預覽結果越準確。")
    add_table(
        doc,
        ["資料", "說明"],
        [
            ["姓名", "報名者姓名，建議放在第一行最前面。"],
            ["課程/場次名稱", "建議包含完整日期，例如 2026/6/13。"],
            ["總金額", "例如 $10600。"],
            ["已付金額", "分期、訂金或部分付款時使用。"],
            ["匯款末五碼或刷卡末四碼", "用來辨識付款來源。"],
            ["付款日期", "例如 3/10 匯款；工具會依課程年份或目前年份推算。"],
            ["備註", "原始文字會整理到 Excel 備註欄。"],
        ],
        [Inches(1.8), Inches(4.7)],
    )
    sample = doc.add_paragraph()
    sample_run = sample.add_run(
        "範例：\n王小明 2026/6/13-6/14 A課程-B進階班 總金額 $10600\n銀行帳號末五碼 40289\n已付金額 $4600\n-\n3/10 匯款"
    )
    set_run_font(sample_run)

    add_heading(doc, "7. Excel 寫入欄位對照")
    add_table(
        doc,
        ["Excel 欄位", "寫入內容"],
        [
            ["A", "月份公式，依付款日期產生 yyyy-mm。"],
            ["B", "付款日期。"],
            ["C", "付款方式，例如匯款或刷卡。"],
            ["D", "匯款末五碼或刷卡末四碼。"],
            ["E", "已付金額。"],
            ["F", "姓名。"],
            ["G", "課程/場次名稱。"],
            ["H", "課程總金額。"],
            ["K", "課程月份，優先從課程日期推算。"],
            ["L", "退款註記，工具保持空白。"],
            ["N", "備註。"],
        ],
        [Inches(1.4), Inches(5.1)],
    )

    add_heading(doc, "8. 分期或多項目資料")
    add_paragraph(doc, "如果文字中包含多個課程項目或分期資訊，工具可能會拆成多筆 Excel 資料列。")
    add_bullet(doc, "一般單筆付款：通常產生 1 筆資料。")
    add_bullet(doc, "多項目付款：可能依項目拆成多筆資料。")
    add_bullet(doc, "分期或訂金：第一筆會帶入已付金額，其餘筆數可能保留空白付款金額。")
    add_paragraph(doc, "寫入前請一定檢查右側預覽，確認筆數和金額符合你的期待。")

    add_heading(doc, "9. 常見問題與處理方式")
    add_table(
        doc,
        ["狀況", "可能原因", "處理方式"],
        [
            ["寫入按鈕不能按", "尚未選 Excel，或預覽資料有錯。", "先選 .xlsx，再補齊左側文字。"],
            ["提示缺少姓名", "文字中沒有明確姓名。", "在第一行最前面放入姓名。"],
            ["提示缺少付款日期", "沒有可辨識的匯款日期。", "補上類似 3/10 匯款 的文字。"],
            ["提示缺少付款代碼", "沒有末五碼或末四碼。", "補上匯款末五碼或刷卡末四碼。"],
            ["寫入失敗", "Excel 檔案正在開啟或沒有權限。", "關閉 Excel 後再執行，必要時另存新檔。"],
            ["課程月份不如預期", "課程名稱中沒有完整日期。", "在課程名稱放入 2026/6/13 這類完整日期。"],
        ],
        [Inches(1.6), Inches(2.15), Inches(2.75)],
    )

    add_heading(doc, "10. 安全使用建議")
    for text in [
        "正式資料第一次使用前，先複製一份 Excel 備份。",
        "每次寫入前先看右側預覽，不要只依賴原始文字。",
        "同一份資料重複按寫入，Excel 會重複新增資料列。",
        "寫入完成後打開 Excel，檢查第 3 列起新增的資料是否正確。",
        "遇到 Windows 不明來源提醒時，只信任來自專案負責人或公司內部可信來源的檔案。",
    ]:
        add_bullet(doc, text)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Excel 報名資料自動填寫工具 使用說明")
    set_run_font(footer_run, size=9, color="777777")

    doc.core_properties.title = "Excel 報名資料自動填寫工具 使用說明"
    doc.core_properties.subject = "User guide for Excel registration auto-fill tool"
    doc.core_properties.author = "PinHaoLin"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
